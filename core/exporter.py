from database import SessionLocal, SupplyRequest, RequestItem, Item, Stock
import os
from datetime import datetime
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Inches, Pt

EXPORT_FILE = "CLEANED_SUPPLY_INVENTORY.xlsx"

def export_to_excel(start_date=None, end_date=None, area=None, shift=None, only_pending=False):
    from datetime import datetime
    from database import parse_frequency, Department
    
    with SessionLocal() as session:
        # 1. Base query for all relevant request items
        query = session.query(RequestItem).options(
            joinedload(RequestItem.item),
            joinedload(RequestItem.supply_request).joinedload(SupplyRequest.employee),
            joinedload(RequestItem.supply_request).joinedload(SupplyRequest.department)
        ).join(SupplyRequest).join(Department)

        # 2. Sequential filters
        if start_date:
            query = query.filter(SupplyRequest.request_date >= start_date)
        if end_date:
            query = query.filter(SupplyRequest.request_date <= end_date)
        if area:
            query = query.filter(Department.area_name == area)
        if shift:
            query = query.filter(Department.shift == shift)

        all_results = query.all()

        if not all_results:
            print("No data found matching filters.")
            return None

        # Flatten the relational data into a list of dictionaries
        data = []
        now = datetime.now()
        
        for req in all_results:
            is_late = req.supply_request.status == "PENDING"
            status = "🔴 PENDING/LATE" if is_late else "🟢 OK"
            
            if only_pending and not is_late:
                continue

            row = {
                "Status": status,
                "Request Date": req.supply_request.request_date.strftime("%Y-%m-%d"),
                "Employee Name": req.supply_request.employee.name,
                "Employee Role": req.supply_request.employee.role,
                "Department Area": req.supply_request.department.area_name,
                "Shift": req.supply_request.department.shift,
                "Supervisor": req.supply_request.department.supervisor,
                "Item Requested": req.item.name,
                "Quantity": req.quantity,
                "Frequency": req.frequency or "N/A"
            }
            data.append(row)

        if not data:
            return None

        df = pd.DataFrame(data)
        
        filename = f"FILTERED_REPORT_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name="Supply Report", index=False)
            worksheet = writer.sheets["Supply Report"]
            for idx, col in enumerate(df):
                series = df[col]
                max_len = max(series.astype(str).map(len).max(), len(str(series.name))) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = max_len

        return filename

def generate_inventory_checklist(data_rows, location_name="ALL"):
    """
    Generates a professional Excel checklist for physical inventory checking.
    data_rows: List of dicts with keys ['Item', 'Threshold', 'Actual', 'Unit', 'Location']
    """
    import pandas as pd
    from openpyxl import Workbook
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Inventory Checklist"
    
    # Styles
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
    center_align = Alignment(horizontal="center", vertical="center")
    thick_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    # Define columns: Added "Checked" column for manual checking
    headers = ["Item Name", "Current Stock", "Unit", "Location", "Threshold", "Checked [ ]"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thick_border
        
    # Column Widths
    ws.column_dimensions['A'].width = 35
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 12
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 15 # Checked column
    
    # Populate Data
    for row_idx, data in enumerate(data_rows, 2):
        ws.cell(row=row_idx, column=1, value=data.get("Item")).border = thick_border
        ws.cell(row=row_idx, column=2, value=data.get("Actual")).border = thick_border
        ws.cell(row=row_idx, column=3, value=data.get("Unit")).border = thick_border
        ws.cell(row=row_idx, column=4, value=data.get("Location")).border = thick_border
        ws.cell(row=row_idx, column=5, value=data.get("Threshold")).border = thick_border
        ws.cell(row=row_idx, column=6, value="").border = thick_border # Blank for checking
        
    # Signatures
    sig_row = len(data_rows) + 4
    ws.cell(row=sig_row, column=1, value="Prepared by: ___________________").font = Font(italic=True)
    ws.cell(row=sig_row, column=4, value="Approved by: ___________________").font = Font(italic=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"INVENTORY_CHECKLIST_{location_name.replace(' ', '_')}_{timestamp}.xlsx"
    wb.save(filename)
    return filename

def generate_stock_confirmation_word(data_rows, location_name="ALL"):
    """
    Generates a formal Word document confirming stock levels.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('CENTRO SERVICES COOPERATIVE', 0)
    title.alignment = 1 # Center
    
    # Info
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("-" * 80)
    
    # Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item Name'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Unit'
    hdr_cells[3].text = 'Location'
    hdr_cells[4].text = 'Price'
    hdr_cells[5].text = 'Total'
    
    # Bold Headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Rows
    grand_total = 0.0
    for data in data_rows:
        price = data.get("Price", 0.0)
        qty = data.get("Actual", 0.0)
        total = price * qty
        grand_total += total
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(data.get("Item"))
        row_cells[1].text = f"{qty:.2f}"
        row_cells[2].text = str(data.get("Unit"))
        row_cells[3].text = str(data.get("Location"))
        row_cells[4].text = f"P{price:,.2f}" if price > 0 else "-"
        row_cells[5].text = f"P{total:,.2f}" if total > 0 else "-"
    
    # Grand Total Row
    if grand_total > 0:
        total_row = table.add_row().cells
        total_row[0].text = "GRAND TOTAL"
        # Merge cells for "GRAND TOTAL" label
        total_row[0].merge(total_row[4])
        total_row[0].paragraphs[0].runs[0].bold = True
        total_row[5].text = f"P{grand_total:,.2f}"
        total_row[5].paragraphs[0].runs[0].bold = True
    
    # Spacing before signature
    doc.add_paragraph("\n\n")
    
    # Signatures
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = True
    cells = sig_table.rows[0].cells
    cells[0].text = "Prepared by:"
    cells[1].text = "Approved by:"
    
    cells = sig_table.rows[1].cells
    cells[0].text = "_________________________"
    cells[1].text = "_________________________"
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"STOCK_CONFIRMATION_{location_name.replace(' ', '_')}_{timestamp}.docx"
    doc.save(filename)
    return filename

if __name__ == "__main__":
    export_to_excel()
